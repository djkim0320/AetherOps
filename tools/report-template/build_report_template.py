from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

INK = "18222C"
DEEP_BLUE = "16334B"
MUTED_BLUE = "486578"
ACCENT = "55D8AE"
ACCENT_DARK = "257862"
ACCENT_PALE = "EAF9F4"
CYAN_PALE = "EAF7FC"
GRAY_050 = "F7F9FA"
GRAY_100 = "EEF2F4"
GRAY_300 = "CBD5DA"
GRAY_500 = "71808A"
WHITE = "FFFFFF"
CAUTION = "8A6500"
CAUTION_PALE = "FFF7DA"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run,
    *,
    latin: str = "Calibri",
    east_asia: str = "Malgun Gothic",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Malgun Gothic")


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, *, top: int = 90, start: int = 130, bottom: int = 90, end: int = 130) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, color: str = GRAY_300, size: int = 6, outer: bool = True) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    edge_names = ("top", "left", "bottom", "right", "insideH", "insideV")
    for edge_name in edge_names:
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        is_outer_edge = edge_name in {"top", "left", "bottom", "right"}
        edge.set(qn("w:val"), "single" if outer or not is_outer_edge else "nil")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_spacing(table, before: float = 5, after: float = 9) -> None:
    previous = table._element.getprevious()
    if previous is not None and previous.tag == qn("w:p"):
        ppr = previous.get_or_add_pPr()
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            ppr.append(spacing)
        spacing.set(qn("w:after"), str(int(before * 20)))
    following = table._element.getnext()
    if following is not None and following.tag == qn("w:p"):
        ppr = following.get_or_add_pPr()
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            ppr.append(spacing)
        spacing.set(qn("w:before"), str(int(after * 20)))


def style_cell_text(cell, *, size: float = 9.5, color: str = INK, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08
        for run in paragraph.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    header_fill: str = ACCENT_PALE,
    compact: bool = False,
) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        cell.text = label
        set_cell_shading(cell, header_fill)
        style_cell_text(cell, size=9.2 if compact else 9.6, color=DEEP_BLUE, bold=True)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = value
            set_cell_shading(cell, WHITE)
            style_cell_text(cell, size=9.0 if compact else 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    set_table_spacing(table)
    return table


def add_field(run, instruction: str, fallback: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def paragraph_border(paragraph, *, side: str, color: str, size: int = 8, space: int = 4) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    pbdr.append(edge)


def paragraph_box(paragraph, *, fill: str, border: str, size: int = 10) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shading = ppr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        ppr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)

    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "7")
        edge.set(qn("w:color"), border)
        pbdr.append(edge)

    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:left"), "180")
    ind.set(qn("w:right"), "180")


def configure_list_styles(doc: Document) -> None:
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style, size=11, color=INK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    # Patch the built-in numbering definitions, rather than relying only on
    # paragraph indentation, so Word and LibreOffice use the exact preset
    # marker, hanging indent, tab stop, spacing, and line-height values.
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        for level in abstract.findall(qn("w:lvl")):
            style_ref = level.find(qn("w:pStyle"))
            if style_ref is None or style_ref.get(qn("w:val")) not in {"ListBullet", "ListNumber"}:
                continue
            ppr = level.find(qn("w:pPr"))
            if ppr is None:
                ppr = OxmlElement("w:pPr")
                level.append(ppr)
            tabs = ppr.find(qn("w:tabs"))
            if tabs is None:
                tabs = OxmlElement("w:tabs")
                ppr.append(tabs)
            for child in list(tabs):
                tabs.remove(child)
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), "720")
            tabs.append(tab)
            ind = ppr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                ppr.append(ind)
            ind.set(qn("w:left"), "720")
            ind.set(qn("w:hanging"), "360")
            spacing = ppr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                ppr.append(spacing)
            spacing.set(qn("w:after"), "160")
            spacing.set(qn("w:line"), "280")
            spacing.set(qn("w:lineRule"), "auto")


def add_list_item(doc: Document, text: str, list_style: str, *, bold_prefix: str | None = None) -> object:
    paragraph = doc.add_paragraph(style=list_style)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=11, color=INK, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=INK)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1) -> object:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(paragraph)
    return paragraph


def add_instruction(doc: Document, text: str, *, after: float = 8) -> object:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, color=GRAY_500, italic=True)
    return paragraph


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> object:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=11, color=INK, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=INK)
    return paragraph


def add_callout(doc: Document, label: str, body: str, *, fill: str = ACCENT_PALE, accent: str = ACCENT_DARK) -> object:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(11)
    paragraph.paragraph_format.line_spacing = 1.12
    paragraph_box(paragraph, fill=fill, border=accent)
    label_run = paragraph.add_run(label.upper())
    set_run_font(label_run, size=9, color=accent, bold=True)
    label_run.add_break()
    body_run = paragraph.add_run(body)
    set_run_font(body_run, size=12.5, color=INK, bold=True)
    return paragraph


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Heading 1": (16, DEEP_BLUE, 16, 8),
        "Heading 2": (13, ACCENT_DARK, 12, 6),
        "Heading 3": (12, MUTED_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    custom_styles = [
        ("AetherOps Kicker", 9.5, ACCENT_DARK, True, 0, 10),
        ("AetherOps Cover Title", 30, DEEP_BLUE, True, 0, 9),
        ("AetherOps Subtitle", 14, MUTED_BLUE, False, 0, 24),
        ("AetherOps Metadata", 10, GRAY_500, False, 0, 3),
        ("AetherOps Caption", 9, GRAY_500, False, 4, 8),
    ]
    for name, size, color, bold, before, after in custom_styles:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, size=size, color=color, bold=bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.08


def configure_page(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    left = paragraph.add_run("AETHEROPS  /  RESEARCH REPORT")
    set_run_font(left, size=8.5, color=GRAY_500, bold=True)
    paragraph.add_run("\t")
    right = paragraph.add_run("〈프로젝트명〉")
    set_run_font(right, size=8.5, color=GRAY_500)
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    paragraph_border(paragraph, side="bottom", color=GRAY_300, size=5, space=4)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left_run = footer_paragraph.add_run("검증 가능한 근거 · 재현 가능한 해석 · 추적 가능한 지식")
    set_run_font(left_run, size=8, color=GRAY_500)
    footer_paragraph.add_run("\t")
    page_label = footer_paragraph.add_run("PAGE ")
    set_run_font(page_label, size=8, color=GRAY_500, bold=True)
    page_field = footer_paragraph.add_run()
    set_run_font(page_field, size=8, color=GRAY_500, bold=True)
    add_field(page_field, "PAGE", "1")


def add_cover(doc: Document, logo_path: Path) -> None:
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.paragraph_format.space_before = Pt(18)
    logo_paragraph.paragraph_format.space_after = Pt(18)
    run = logo_paragraph.add_run()
    shape = run.add_picture(str(logo_path), width=Inches(0.82))
    docpr = shape._inline.docPr
    docpr.set("descr", "AetherOps logo")

    kicker = doc.add_paragraph("AETHEROPS RESEARCH", style="AetherOps Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = doc.add_paragraph("연구 보고서", style="AetherOps Cover Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "검증 가능한 근거 · 재현 가능한 해석 · 추적 가능한 지식",
        style="AetherOps Subtitle",
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    report_title = doc.add_paragraph()
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_title.paragraph_format.space_after = Pt(22)
    report_title_run = report_title.add_run("〈연구 제목을 한 문장으로 입력〉")
    set_run_font(report_title_run, size=18, color=INK, bold=True)

    metadata = add_table(
        doc,
        ["프로젝트", "연구 실행"],
        [["〈프로젝트명〉", "〈Run ID〉"], ["〈대화 세션〉", "〈완료 일시 / 시간대〉"]],
        [4680, 4680],
        header_fill=GRAY_100,
        compact=True,
    )
    for row in metadata.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    pillars = doc.add_table(rows=1, cols=3)
    set_table_geometry(pillars, [3120, 3120, 3120])
    set_table_borders(pillars, color=ACCENT, size=6)
    set_repeat_table_header(pillars.rows[0])
    for cell, title_text, body_text in zip(
        pillars.rows[0].cells,
        ("근거", "분석", "판정"),
        (
            "원문과 산출물까지 추적",
            "조건·도구·계산을 재현",
            "독립 리뷰와 품질 게이트",
        ),
    ):
        set_cell_shading(cell, ACCENT_PALE)
        set_cell_margins(cell, top=140, start=120, bottom=140, end=120)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = cell.paragraphs[0].add_run(title_text)
        set_run_font(title_run, size=10, color=ACCENT_DARK, bold=True)
        body_p = cell.add_paragraph()
        body_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        body_p.paragraph_format.space_after = Pt(0)
        body_run = body_p.add_run(body_text)
        set_run_font(body_run, size=8.8, color=INK)

    usage = doc.add_paragraph()
    usage.paragraph_format.space_before = Pt(20)
    usage.paragraph_format.space_after = Pt(0)
    usage.alignment = WD_ALIGN_PARAGRAPH.CENTER
    usage_run = usage.add_run(
        "템플릿 안내  ·  꺾쇠 괄호 안의 안내 문구는 실제 연구 내용으로 교체하고, 해당 없는 조건부 장은 제거합니다."
    )
    set_run_font(usage_run, size=8.5, color=GRAY_500, italic=True)


def add_summary_page(doc: Document, bullet_id: int) -> None:
    add_heading(doc, "한눈에 보는 결론")
    add_instruction(doc, "의사결정자가 첫 화면만 읽어도 연구 결론, 핵심 근거의 강도, 즉시 취해야 할 조치가 명확히 드러나도록 작성합니다.")
    add_callout(
        doc,
        "핵심 결론",
        "〈연구 질문에 대한 직접적인 답을 2~4문장으로 명확히 제시합니다. 결론을 서두에 배치하고 필수 조건과 예외 사항을 명시합니다.〉",
    )

    add_heading(doc, "핵심 발견", level=2)
    add_list_item(doc, "발견 1: 〈가장 결정적인 정량·정성 연구 결과와 문장 밀착 인용 표식 [1]〉", bullet_id, bold_prefix="발견 1:")
    add_list_item(doc, "발견 2: 〈두 번째 핵심 결과 및 이를 뒷받침하는 교차 검증 근거 [2]〉", bullet_id, bold_prefix="발견 2:")
    add_list_item(doc, "발견 3: 〈주요 반대 근거, 해석 경계 조건 또는 예상 외 수치 결과 [3]〉", bullet_id, bold_prefix="발견 3:")

    add_heading(doc, "판단과 다음 행동", level=2)
    add_table(
        doc,
        ["구분", "판단", "근거 수준", "권고 행동"],
        [
            ["핵심 결정", "〈채택 / 조건부 진행 / 보류 / 기각〉", "〈높음 / 중간 / 낮음〉", "〈즉시 실행할 구체적 조치〉"],
            ["주요 위험", "〈의사결정에 영향을 미치는 핵심 리스크〉", "〈검증됨 / 추가 확인 필요〉", "〈사전 완화 대책〉"],
            ["후속 연구", "〈미해결 질문 또는 확장 검증 과제〉", "〈필요 자료 및 도구〉", "〈실행 조건 및 마일스톤〉"],
        ],
        [1440, 2880, 1800, 3240],
        compact=True,
    )

    add_heading(doc, "품질 상태", level=2)
    add_table(
        doc,
        ["인용 무결성", "지식 근거", "치명적 오류", "평균 리뷰 점수", "판정"],
        [["〈100%〉", "〈100% / unsupported 0〉", "〈0건〉", "〈4.0 / 5.0 이상〉", "〈통과 / 품질 실패〉"]],
        [1800, 2340, 1620, 1800, 1800],
        header_fill=CYAN_PALE,
        compact=True,
    )


def add_scope_method_page(doc: Document, bullet_id: int, decimal_id: int) -> None:
    add_heading(doc, "연구 질문과 범위")
    add_heading(doc, "연구 질문", level=2)
    add_body(doc, "〈사용자가 실제로 결정하거나 해결하려는 질문과 목표를 원문의 의미를 정확히 유지하여 서술합니다.〉")
    add_heading(doc, "범위와 제외 범위", level=2)
    add_table(
        doc,
        ["포함 범위", "제외 범위 및 사유"],
        [["〈연구 대상, 기간, 운용 조건, 분석 기법, 데이터셋〉", "〈의도적으로 다루지 않은 대상, 미적용 가정 및 그 사유〉"]],
        [4680, 4680],
        compact=True,
    )
    add_heading(doc, "성공 기준", level=2)
    add_list_item(doc, "〈핵심 연구 질문에 대한 직접적이고 정량적인 해답 제시〉", bullet_id)
    add_list_item(doc, "〈독립 출처 다양성 및 상호 교차 검증성 확보〉", bullet_id)
    add_list_item(doc, "〈공학 시뮬레이션 시 수치 수렴성 및 격자/패널 독립성 입증〉", bullet_id)

    add_heading(doc, "방법과 근거")
    add_instruction(doc, "수행한 절차뿐만 아니라 해당 방법론의 타당성과 데이터 검증 경로를 종합적으로 설명합니다.")
    for text in (
        "계획: 〈연구 가설, 분할 워크스트림, 활용 도구/솔버 및 사전 수락 기준 수립〉",
        "수집: 〈출처 선정 기준, 검색/수집 쿼리, 운용 및 공학 계산 조건 설정〉",
        "병합: 〈다양한 출처 및 계산 결과의 비교 분석, 데이터 충돌 해소 및 통합〉",
        "리뷰: 〈독립 품질 게이트 검토, 인용/지식 무결성 검증 및 피드백 반영〉",
    ):
        add_list_item(doc, text, decimal_id)

    add_heading(doc, "근거 구성", level=2)
    add_table(
        doc,
        ["근거 묶음", "역할", "출처 및 도구", "검증 상태"],
        [
            ["〈Workstream A〉", "〈문헌·기술·시장 조사 분석〉", "〈공개 학술 데이터베이스〉", "〈CAS readback 완료〉"],
            ["〈Workstream B〉", "〈대조군 비교 및 반증 분석〉", "〈공식 기술 사양서/특허〉", "〈CAS readback 완료〉"],
            ["〈공학 검증〉", "〈수치해석 및 민감도 검증〉", "〈SU2 / XFOIL 등 솔버〉", "〈Receipt 검증 완료〉"],
        ],
        [1800, 2520, 2520, 2520],
        compact=True,
    )


def add_findings_engineering_page(doc: Document, bullet_id: int) -> None:
    add_heading(doc, "핵심 결과")
    add_instruction(doc, "각 결과는 [주장 → 근거 → 시사점 → 예외/경계조건]의 논리 흐름으로 기술하며, 인용 표식을 문장에 밀착 배치합니다.")

    add_heading(doc, "발견 A — 〈핵심 주장을 간결한 제목으로 표현〉", level=2)
    add_body(doc, "〈주요 결과를 설명합니다. 모든 정량 지표에는 운용 조건, 물리 단위, 비교 기준을 명시하고 인용 표식 [1]을 배치합니다.〉")
    add_callout(doc, "시사점 및 중요성", "〈이 결과가 사용자의 최종 의사결정, 제품 설계 또는 후속 공학 실험에 미치는 실질적 영향을 서술합니다.〉", fill=CYAN_PALE, accent=MUTED_BLUE)

    add_heading(doc, "발견 B — 〈반대 근거 또는 조건부 결과 분석〉", level=2)
    add_body(doc, "〈결론을 제한하거나 특정 조건에서 결과가 달라질 수 있는 반대 근거 및 불일치 요인을 객관적으로 분석합니다 [2].〉")

    add_heading(doc, "비교 결과", level=2)
    add_table(
        doc,
        ["대안 / 조건", "핵심 성능 지표", "측정 / 계산값", "비교 평가 및 해석", "근거"],
        [
            ["〈기준선 (Baseline)〉", "〈지표 / 단위〉", "〈값〉", "〈기준 성능〉", "〈[1]〉"],
            ["〈후보 A (Candidate A)〉", "〈지표 / 단위〉", "〈값〉", "〈개선도 및 우수 요인〉", "〈[2]〉"],
            ["〈후보 B (Candidate B)〉", "〈지표 / 단위〉", "〈값〉", "〈한계점 및 평가〉", "〈[3]〉"],
        ],
        [1800, 1800, 1440, 2880, 1440],
        compact=True,
    )

    add_page_break(doc)
    add_heading(doc, "공학 해석 (조건부)")
    add_instruction(doc, "독립 공학 시뮬레이션 및 수치 계산이 포함되지 않은 일반 연구의 경우 이 장 전체를 생략합니다.")
    add_heading(doc, "해석 계약과 조건", level=2)
    add_table(
        doc,
        ["항목", "해석 설정 및 형상", "근거 및 Receipt"],
        [
            ["형상 및 기하학적 모델", "〈형상 ID, 파라미터, 물성치〉", "〈artifact / CAS pointer〉"],
            ["경계 및 운용 조건", "〈Mach, Re, AoA, 유동 조건 등〉", "〈artifact / CAS pointer〉"],
            ["수치 해석 파라미터", "〈격자 밀도, 난류 모델, 반복 수, 잔차 기준〉", "〈config / receipt〉"],
            ["민감도 및 독립 검증", "〈격자/패널 세분화, 조건 변화 분석〉", "〈비교 검증 receipt〉"],
        ],
        [2160, 4320, 2880],
        compact=True,
    )
    add_heading(doc, "결과와 수렴성", level=2)
    add_body(doc, "〈수치 해석 결과를 물리적 메커니즘과 함께 심층 해석합니다. 계산 완료와 수치적 수렴, 과학적 결론 확정을 엄격히 구분하고, 격자 의존성 및 민감도 문제를 명시합니다.〉")
    add_instruction(doc, "그림 1. 〈비교 그래프 또는 수렴 이력 곡선〉 — 좌표축, 단위, 기준선, 불확실성 오차 범위를 캡션에 명시합니다.")


def add_conclusion_page(doc: Document, bullet_id: int) -> None:
    add_heading(doc, "결론 및 권고")
    add_callout(
        doc,
        "최종 권고",
        "〈가장 타당한 전략적 대안과 해당 권고가 유효한 전제 조건을 명확히 제시합니다. 추가 검증이 필요할 경우 부족한 요소를 직접 밝힙니다.〉",
    )
    add_heading(doc, "실행 우선순위 및 로드맵", level=2)
    add_table(
        doc,
        ["우선순위", "실행 과제", "완료 조건 및 목표치", "의존성 및 리스크"],
        [
            ["1 (최우선)", "〈즉시 착수할 핵심 조치〉", "〈정량적으로 검증 가능한 완료 기준〉", "〈선행 의존 요소〉"],
            ["2 (차선)", "〈후속 설계 및 프로세스 조치〉", "〈정량적으로 검증 가능한 완료 기준〉", "〈잠재 리스크 및 완화책〉"],
            ["3 (모니터링)", "〈지속 검증 또는 모니터링〉", "〈프로젝트 중단/전환 기준〉", "〈필요 자원 및 승인〉"],
        ],
        [1200, 3000, 2760, 2400],
        compact=True,
    )

    add_heading(doc, "한계와 불확실성")
    add_list_item(doc, "자료 한계: 〈표본 크기, 수집 기간, 데이터 접근 한계 또는 출처의 잠재적 편향〉", bullet_id, bold_prefix="자료 한계:")
    add_list_item(doc, "방법 한계: 〈적용된 물리 모델, 단순화 가정, 해석기 알고리즘의 유효 한계 영역〉", bullet_id, bold_prefix="방법 한계:")
    add_list_item(doc, "결론 민감도: 〈주요 환경 변수 또는 가정 변화 시 결론의 변동 민감도〉", bullet_id, bold_prefix="결론 민감도:")
    add_list_item(doc, "남은 질문: 〈현재 데이터로 확답할 수 없으며 추가 연구가 필요한 핵심 과제〉", bullet_id, bold_prefix="남은 질문:")

    add_heading(doc, "출처")
    add_instruction(doc, "본문의 인용 표식과 1:1로 정확히 대응하도록 작성하며, 저자/기관, 제목, 발행처, 일자, URL/식별자를 표준 양식으로 제공합니다.")
    add_body(doc, "[1] 〈저자/기관〉. 〈자료 제목〉. 〈발행처〉, 〈발행일자〉. 〈URL 또는 고유 식별자〉.")
    add_body(doc, "[2] 〈저자/기관〉. 〈자료 제목〉. 〈발행처〉, 〈발행일자〉. 〈URL 또는 고유 식별자〉.")
    add_body(doc, "[3] 〈공학 계산 receipt 또는 검증된 내부 산출물의 사람이 읽을 수 있는 명칭〉.")


def add_audit_page(doc: Document) -> None:
    add_heading(doc, "재현성 및 감사 부록")
    add_heading(doc, "연구 실행 정보", level=2)
    add_table(
        doc,
        ["항목", "검증된 값"],
        [
            ["연구 프로필", "〈research_profile_version〉"],
            ["검색 프로필", "〈hybrid_graph_v1〉"],
            ["지식 generation", "〈knowledge_generation_id〉"],
            ["PLAN / MERGE / REVIEW", "〈모델 · 추론 수준 · service tier〉"],
            ["COLLECT", "〈모델 · 추론 수준 · collector 수〉"],
            ["도구 및 솔버", "〈도구명 · 버전 · receipt ID〉"],
        ],
        [2520, 6840],
        compact=True,
    )

    add_heading(doc, "지식그래프 반영", level=2)
    add_table(
        doc,
        ["개체 (Entity)", "주장 (Assertion)", "추론 (Inference)", "충돌 (Conflict)", "RDF Snapshot"],
        [["〈수량〉", "〈수량 / unsupported 0〉", "〈수량 / proof chain〉", "〈0건 또는 검토 항목〉", "〈SHA-256 / Triple 수〉"]],
        [1600, 2100, 1860, 1800, 2000],
        header_fill=ACCENT_PALE,
        compact=True,
    )

    add_heading(doc, "품질 리뷰 판정", level=2)
    add_table(
        doc,
        ["평가 축", "점수", "판정 근거 요약"],
        [
            ["과업 충족도", "〈1–5〉", "〈핵심 평가 근거〉"],
            ["주장 근거성", "〈1–5〉", "〈핵심 평가 근거〉"],
            ["출처 품질 및 다양성", "〈1–5〉", "〈핵심 평가 근거〉"],
            ["반대 근거 및 완전성", "〈1–5〉", "〈핵심 평가 근거〉"],
            ["추론 및 불확실성", "〈1–5〉", "〈핵심 평가 근거〉"],
            ["명료성 및 재현성", "〈1–5〉", "〈핵심 평가 근거〉"],
        ],
        [2520, 1200, 5640],
        compact=True,
    )
    add_callout(
        doc,
        "품질 게이트 통과 규칙",
        "인용 무결성 100%, 지식 근거 무결성 100%, unsupported assertion 0건, 치명적 오류 0건, 모든 평가 축 3점 이상, 전체 산술 평균 4.0점 이상.",
        fill=GRAY_050,
        accent=DEEP_BLUE,
    )

    add_heading(doc, "산출물 무결성", level=2)
    add_instruction(doc, "사용자 보고서 본문에는 사람이 이해하기 쉬운 명칭을 사용하며, 해시 및 내부 ID는 감사 부록에 격리 보관합니다.")
    add_table(
        doc,
        ["산출물 명칭", "종류", "SHA-256 / Artifact ID"],
        [
            ["〈최종 연구 보고서〉", "보고서 (Report)", "〈hash / artifact ID〉"],
            ["〈근거 자료 묶음〉", "근거 (Evidence)", "〈hash / artifact ID〉"],
            ["〈공학 해석 결과〉", "공학 (Engineering)", "〈hash / artifact ID〉"],
            ["〈품질 검토 보고서〉", "리뷰 (Verdict)", "〈hash / artifact ID〉"],
        ],
        [3600, 1600, 4160],
        compact=True,
    )

    add_heading(doc, "발행 전 최종 점검표", level=2)
    for item in (
        "본문 인용 표식과 citation manifest 및 원본 출처가 1:1로 완전하게 대응한다.",
        "모든 수치 데이터에 표준 단위(SI)와 운용/해석 경계 조건이 함께 명시되어 있다.",
        "계산 완료, 수치적 수렴, 과학적 결론 확정을 엄격히 구분하여 서술하였다.",
        "사용자 친화적 명칭을 사용하고 내부 시스템 ID/해시는 감사 부록으로 분리하였다.",
        "한계점, 반대 근거, 미해결 질문을 축소하거나 은폐하지 않고 투명하게 기술하였다.",
    ):
        add_list_item(doc, item, "List Bullet")


def build_document(output_path: Path, logo_path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    configure_list_styles(doc)
    bullet_id = "List Bullet"
    decimal_id = "List Number"

    props = doc.core_properties
    props.title = "AetherOps 연구 보고서 템플릿"
    props.subject = "근거 기반 일반·공학 연구 보고서"
    props.author = "AetherOps"
    props.keywords = "AetherOps, research, report, evidence, engineering, knowledge graph"
    props.comments = "AetherOps report artifact template v1"

    add_cover(doc, logo_path)
    add_page_break(doc)
    add_summary_page(doc, bullet_id)
    add_page_break(doc)
    add_scope_method_page(doc, bullet_id, decimal_id)
    add_page_break(doc)
    add_findings_engineering_page(doc, bullet_id)
    add_page_break(doc)
    add_conclusion_page(doc, bullet_id)
    add_page_break(doc)
    add_audit_page(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the AetherOps research report DOCX template.")
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--logo", required=True, type=Path)
    args = parser.parse_args()
    if not args.logo.is_file():
        raise SystemExit(f"logo does not exist: {args.logo}")
    build_document(args.output.resolve(), args.logo.resolve())


if __name__ == "__main__":
    main()
